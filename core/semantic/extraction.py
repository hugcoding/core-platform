from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Callable, Iterable


SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


def _docx_text(path: Path) -> tuple[str, int]:
    from docx import Document

    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks), 0


def _pdf_text(path: Path) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise PermissionError("password-protected PDF") from exc
        if not unlocked:
            raise PermissionError("password-protected PDF")
    return "\n".join(page.extract_text() or "" for page in reader.pages), len(reader.pages)


def extract_document(
    path: Path,
    *,
    docx_loader: Callable[[Path], tuple[str, int]] = _docx_text,
    pdf_loader: Callable[[Path], tuple[str, int]] = _pdf_text,
) -> tuple[str, int]:
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"unsupported extension: {extension or '[none]'}")
    if not path.is_file():
        raise FileNotFoundError(path)
    return (docx_loader if extension == ".docx" else pdf_loader)(path)


def extract_statistics(
    path: Path,
    *,
    docx_loader: Callable[[Path], tuple[str, int]] = _docx_text,
    pdf_loader: Callable[[Path], tuple[str, int]] = _pdf_text,
) -> dict[str, Any]:
    extension = path.suffix.lower()
    text, pages = extract_document(
        path,
        docx_loader=docx_loader,
        pdf_loader=pdf_loader,
    )
    normalized = " ".join(text.split())
    return {
        "extension": extension.removeprefix("."),
        "size_bytes": path.stat().st_size,
        "characters": len(normalized),
        "words": len(normalized.split()) if normalized else 0,
        "pages": pages or None,
        "has_extractable_text": bool(normalized),
    }


def run_manifest(
    manifest_path: Path,
    *,
    extractor: Callable[[Path], dict[str, Any]] = extract_statistics,
) -> Iterable[dict[str, Any]]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest.get("processing") != "local_only":
        raise ValueError("pilot manifest must require local_only processing")
    if manifest.get("embedding_enabled") is not False:
        raise ValueError("embeddings must be disabled for the extraction pilot")
    if manifest.get("external_ai_enabled") is not False:
        raise ValueError("external AI must be disabled for the extraction pilot")
    if manifest.get("database_writes_enabled") is not False:
        raise ValueError("database writes must be disabled for the extraction pilot")

    for item in manifest["files"]:
        if item["approval"] != "approved":
            yield {
                "file_id": item["file_id"],
                "status": "skipped",
                "reason": f"approval={item['approval']}",
            }
            continue

        try:
            statistics = extractor(Path(item["path"]))
        except Exception as exc:
            yield {
                "file_id": item["file_id"],
                "status": "error",
                "error_type": type(exc).__name__,
                "reason": str(exc),
            }
            continue

        yield {"file_id": item["file_id"], "status": "extracted", **statistics}


def summarize(results: list[dict[str, Any]]) -> dict[str, Any]:
    extracted = [result for result in results if result["status"] == "extracted"]
    no_text = [result for result in extracted if not result.get("has_extractable_text")]
    return {
        "status": "summary",
        "documents": len(results),
        "extracted": len(extracted),
        "extractable_text": len(extracted) - len(no_text),
        "needs_ocr": sum(
            result.get("extension") == "pdf" for result in no_text
        ),
        "no_text": len(no_text),
        "password_protected": sum(
            result.get("error_type") == "PermissionError" for result in results
        ),
        "skipped": sum(result["status"] == "skipped" for result in results),
        "errors": sum(result["status"] == "error" for result in results),
        "characters": sum(int(result.get("characters") or 0) for result in extracted),
        "words": sum(int(result.get("words") or 0) for result in extracted),
        "pages": sum(int(result.get("pages") or 0) for result in extracted),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(prog="semantic-pilot-extract")
    parser.add_argument(
        "--manifest",
        default="/app/project/pilots/scrum-57-documents-v1.json",
    )
    args = parser.parse_args(argv)

    results = list(run_manifest(Path(args.manifest)))
    for result in results:
        print(json.dumps(result, ensure_ascii=False))
    summary = summarize(results)
    print(json.dumps(summary, ensure_ascii=False))
    return 1 if summary["errors"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
