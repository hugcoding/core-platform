try:
    from docx import Document
except Exception:
    Document = None


def render_docx(md_path, docx_path):
    if Document is None:
        return False

    text = md_path.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    in_code = False

    for line in text.splitlines():
        if line.startswith("```"):
            in_code = not in_code
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip() == "---":
            doc.add_page_break()
        elif in_code:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return True
